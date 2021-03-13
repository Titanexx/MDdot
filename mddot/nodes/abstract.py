import re

from logger import logger
import helpers

from anytree import NodeMixin
from docxtpl import RichText, InlineImage
from docx.shared import Length
from PIL import Image

class AbstractNode( NodeMixin):
	template = None
	docxTemplate = None
	mdFilename = None
	
	styles = {}
	templateStyles = {}
	keyTplCache = {}
	
	images = {}
	tables = {}

	rawxml = ""
	
	classByToken = {}
	
	def __init_subclass__(cls, tokenClass=None, **kwargs):
		if tokenClass:
			super().__init_subclass__(**kwargs)
			logger.verbose("Register %s for %s" % (cls, tokenClass))
			if cls.styles:
				AbstractNode.styles |= cls.styles
				cls.styles = AbstractNode.styles
			if tokenClass not in cls.classByToken: 
				AbstractNode.classByToken[tokenClass] = [cls]
			else: 
				AbstractNode.classByToken[tokenClass].append(cls)

	@classmethod
	@staticmethod
	def createFromToken(token,parent):
		tokenType = type(token)
		good_class = AbstractNode.classByToken['none'][0]
		if not tokenType in AbstractNode.classByToken:
			logger.warning("Ignore %s because it's not implemented." % tokenType)
		else:
			for c in AbstractNode.classByToken[tokenType]:
				if c.condition(token,parent):
					good_class = c
				else:
					logger.spam("%s condition isn't good." % c)
			logger.debug("Generate %s for %s with parent %s" % (good_class,tokenType,parent))

		return good_class(token,parent)

	@staticmethod
	def condition(block,parent):
		return True
	
	def __init__(self, block=None, parent=None, children=[]):
		super().__init__()
		self._id = ""
		self._block = block
		self.parent = parent
		self.children = children

	def __str__(self):
		return "<%s,id=%s>" % (self.__class__.__name__,self.id)

	def _checkStyles(self, log=True):
		for k,s in self.styles.items():
			for st in self.docxTemplate.styles:
				if st.style_id == s:
					AbstractNode.templateStyles[k] = st
					break
			else:
				if log:
					logger.warning("Can't found '%s' style inside template." % s)

	def _buildKeyTplCache(self):
		keyRegex = re.compile(r'{{ (.*?)(\.(xml))? }}')
		forRegex = re.compile(r'\[[a-z]+\]')
		paragraphs = []
		for p in self.docxTemplate.paragraphs:
			paragraphs.append(p)
		for t in self.docxTemplate.tables:
			for r in t.rows:
				for c in r.cells:
					for p in c.paragraphs:
						if p not in paragraphs:
							# Avoid duplicate, should not append
							paragraphs.append(p)

		for p in paragraphs:
			key = keyRegex.search(p.text)
			if key:
				key = forRegex.sub(".*",key[1])
				if key[-1] == '.':
					key += '.'
				pPr = helpers.getpPr(p)
				logger.spam("Add (key , p , pPr) in cache : (%s , %s , %s)" % (key,p,pPr))
				self.keyTplCache[key] = (p,pPr)

	def getTplById(self, id):
		if id in self.keyTplCache:
			return self.keyTplCache[id]
		else:
			# handle for loop
			# Need jinja declaration as list not with items inside template 
			for key in self.keyTplCache:
				if helpers.compareKeys(id,key):
					return self.keyTplCache[key]
			return ("","")

	def getStyleById(self, id=None):
		if not id:
			id = self.id
		return self.getTplById(id)[1]

	def getPById(self, id=None):
		if not id:
			id = self.id
		return self.getTplById(id)[0]

	def finalizeImages(self,tpl):
		self.setTemplate(tpl,log=False)

		for id,img in self.images.items():
			p = self.getPById(img['id'])
			im = Image.open(img['src'])

			if "_Body" in str(type(p._parent)):
				section = p._parent._element.xpath(".//following::w:sectPr")[0]
				length=section.page_width - (section.right_margin + section.left_margin) 
			else:			
				rightMargin = helpers.getMarginFromStyle(p._parent._parent.style,'right')
				leftMargin = helpers.getMarginFromStyle(p._parent._parent.style,'left')
				length = Length(p._parent.width) - (rightMargin + leftMargin)

			width, height = im.size
			dpiWidth, dpiHeight = im.info['dpi']
			widthEmus = width / dpiWidth * Length._EMUS_PER_INCH
			heightEmus = height / dpiHeight * Length._EMUS_PER_INCH
			imgWidth = widthEmus
			if widthEmus > length:
				imgWidth = length

			logger.debug("Load image %s inside %s." % (img,tpl))
			img['image'] = InlineImage(tpl,img['src'],imgWidth)

	def setTemplate(self,template,log=True):
		AbstractNode.template = template
		AbstractNode.docxTemplate = template.docx
		self._checkStyles(log)
		self._buildKeyTplCache()

	@property   
	def id(self):
		if not self.parent:
			return ""
		else:
			pid = self.parent.id
			if self._id:
				if pid:
					return "%s.%s" % (pid,self._id)
				else:
					return self._id
			else:
				return pid

	def generate(self):
		logger.warning("%s must implement generate method" % self.__class__)
		return {}
